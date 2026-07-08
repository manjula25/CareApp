#!/usr/bin/env python3
"""Minimal Markdown -> DOCX converter for the CareSync design spec.

Handles the subset of Markdown used in docs/DESIGN_SPEC.md: ATX headings,
fenced code blocks, GitHub-style tables, blockquotes, bullet lists, and
inline `code`/**bold** spans. Not a general Markdown engine.
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(paragraph, text):
    """Render inline **bold** and `code` spans into a paragraph."""
    token = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(part)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell in enumerate(header):
        p = table.rows[0].cells[i].paragraphs[0]
        add_inline(p, cell)
        for run in p.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            val = row[i] if i < len(row) else ""
            add_inline(cells[i].paragraphs[0], val)


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main(md_path, docx_path):
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(6)
            continue

        # Table
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = parse_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        stripped = line.strip()

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            h.text = ""
            add_inline(h, m.group(2))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip(">").strip())
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            content = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
